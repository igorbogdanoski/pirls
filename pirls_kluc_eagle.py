from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_pirls_key():
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)

    title = doc.add_paragraph("УПАТСТВО ЗА ОЦЕНУВАЊЕ (КЛУЧ): Летај, Орле, летај (15 Прашања)")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True

    doc.add_paragraph("Текст: Африканска приказна\n")

    mcq_answers = [
        ("1", "А"), ("3", "А"), ("6", "А"), ("7", "Б"), ("12", "А"), ("14", "В")
    ]

    doc.add_paragraph("Прашања со повеќечлен избор (1 поен за точен одговор):").runs[0].bold = True
    for num, ans in mcq_answers:
        doc.add_paragraph(f"Прашање {num}: {ans}")

    doc.add_paragraph("\nПрашања од отворен тип:").runs[0].bold = True
    
    open_guide = [
        ("2", "Локација", "1 бод: На раб на карпа / на планината."),
        ("4", "Поука", "1 бод: Сакал да го научи да биде кокошка."),
        ("5", "Примери", "2 бода: Оди како кокошка, јаде како кокошка, размислува како кокошка (наведени барем два)."),
        ("8", "Значење", "2 бода: Интерпретација дека неговото место е во висините/слободата, а не во заробеништво."),
        ("9", "Смеа", "1 бод: Бидејќи орелот се вратил кај кокошките и докажал дека фармерот е во право."),
        ("10", "Причини", "2 бода: Да го види сонцето како изгрева; да биде на местото каде што припаѓа (високо)."),
        ("11", "Опис", "1 бод: Розови / со златен сјај."),
        ("13", "Особина", "2 бода: Упорен (не се откажал по првиот неуспех) / Мудар (знаел како да го поттикне орелот)."),
        ("15", "Наслов", "2 бода: Соодветен е бидејќи ја претставува главната порака на охрабрување за остварување на потенцијалот.")
    ]

    for num, name, criteria in open_guide:
        p = doc.add_paragraph()
        p.add_run(f"Прашање {num} ({name}): ").bold = True
        doc.add_paragraph(criteria)

    doc.save("PIRLS_Key_Eagle_15Q.docx")

create_pirls_key()
