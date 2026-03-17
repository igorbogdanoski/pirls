from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_pirls_key_zajak():
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)

    title = doc.add_paragraph("УПАТСТВО ЗА ОЦЕНУВАЊЕ (КЛУЧ): Зајакот го најавува земјотресот (15 Прашања)")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True

    doc.add_paragraph("Текст: Класична приказна\n")

    mcq_answers = [
        ("1", "В"), ("2", "Б"), ("5", "В"), ("6", "В"), ("11", "Б"), ("12", "В")
    ]

    doc.add_paragraph("Прашања со повеќечлен избор (1 поен за точен одговор):").runs[0].bold = True
    for num, ans in mcq_answers:
        doc.add_paragraph(f"Прашање {num}: Точен одговор {ans} (1 бод)")

    doc.add_paragraph("\nПрашања од отворен тип:").runs[0].bold = True
    
    open_guide = [
        ("3", "Зборови за брзина", "1 бод: стрчал, споулавени, потрчав, брзо, бегајте."),
        ("4", "Локација", "1 бод: Сакал да го однесе назад кај неговиот дом / кај дрвото."),
        ("7", "Утеха од лавот", "2 бода: Го нарекол 'братко', му се насмевнал, му објаснил дека сите се плашат."),
        ("8", "Мислење", "2 бода: Да, бидејќи бил трпелив со него и му помогнал да ја надмине паниката."),
        ("9", "Промена на чувства", "3 бода: Почеток (исплашен/земјотрес); Крај (смирен/глупаво/плод)."),
        ("10", "Разлики лав/зајак", "3 бода: Зајакот е паничар (бега без доказ); Лавот е мудар (истражува факти)."),
        ("13", "Согласност", "1 бод: Бидејќи почувствувал дека може да му верува на лавот."),
        ("14", "Споредба гром", "2 бода: Покажува дека илјадници зајаци биле многу гласни додека трчале."),
        ("15", "Паметна постапка", "2 бода: Да, бидејќи само така зајакот можел навистина да се увери во вистината.")
    ]

    for num, name, criteria in open_guide:
        p = doc.add_paragraph(f"\nПрашање {num}: {name}")
        p.runs[0].bold = True
        doc.add_paragraph(criteria)

    doc.save("PIRLS_Kluc_Zajak_Nastavnik.docx")

create_pirls_key_zajak()
