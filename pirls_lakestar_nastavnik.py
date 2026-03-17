from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

def create_lake_star_key():
    doc = Document()
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(12)
    doc.styles['Normal'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    doc.add_heading('УПАТСТВО ЗА ОЦЕНУВАЊЕ (КЛУЧ): Спасувањето на „Езерската ѕвезда“', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER

    mcq = [("1", "В"), ("4", "Г"), ("5", "Б"), ("6", "В"), ("8", "В"), ("12", "В"), ("14", "В")]
    doc.add_paragraph("ПРАШАЊА СО ПОВЕЌЕЧЛЕН ИЗБОР (1 бод)").runs[0].bold = True
    for q, a in mcq:
        doc.add_paragraph(f"Прашање {q}: Точен одговор {a}")

    open_q = [
        ("2", "Што му било потребно", "1 бод: Наведува ДВЕ од следниве: шмиргла (за дрво), (нова сина) боја, трпение."),
        ("3", "Чувствата на Тео", "2 бода: Возбуден/среќен, затоа што отсекогаш сонувал да има своја едрилица."),
        ("7", "Споредба на бродчињата", "2 бода: Наведува две контрастни разлики (пр. На Тео е од дрво / на Бојан од пластика; На Тео на ветер / на Бојан на мотор; Старо наспроти ново)."),
        ("9", "Зошто се пробила лесно", "1 бод: Немало мотор што би се заплеткал во тревите ИЛИ било направено од цврсто дрво."),
        ("10", "Особина на Тео", "2 бода: Вреден/трпелив (го шмирглал дрвото) ИЛИ добар пријател/несебичен (му помогнал на Бојан иако му се смеел)."),
        ("11", "Чувствата на Бојан", "2 бода: На почетокот: потсмев (изгледало како ѓубре/старо). На крајот: восхит/благодарност (сфатил дека е најхрабро откако го спасило неговото)."),
        ("13", "Наслов", "2 бода: Став (Да) + објаснување дека кулминира со херојскиот чин на спасување на пластичното бродче од страна на Езерската ѕвезда.")
    ]

    doc.add_paragraph("\nПРАШАЊА ОД ОТВОРЕН ТИП").runs[0].bold = True
    for q, name, criteria in open_q:
        p = doc.add_paragraph()
        p.add_run(f"Прашање {q} ({name}): ").bold = True
        doc.add_paragraph(f"• {criteria}")

    doc.save("PIRLS_LakeStar_Nastavnik.docx")
    print("Teacher key generated!")

create_lake_star_key()
