from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING

# ==========================================
# 1. ПОДАТОЦИ ЗА ТЕСТОТ (Менувај ги овие за секој нов тест)
# ==========================================

NASLOV_NA_TEKST = "Тајната на стариот часовничар"

# Прашања со повеќечлен избор (MCQ)
# Формат: (Број, Точен одговор, Когнитивен процес)
MCQ_ANSWERS = [
    ("1", "Б", "Пронаоѓање на експлицитно наведени информации"),
    ("3", "В", "Донесување директни заклучоци"),
    ("5", "В", "Интерпретација и интегрирање на идеи"),
    ("6", "Б", "Испитување и вреднување на јазични елементи"),
    ("8", "Б", "Интерпретација и интегрирање на идеи"),
    ("10", "В", "Пронаоѓање на експлицитно наведени информации")
]

# Прашања од отворен тип
# Формат: Речник со прашањето, процесот и критериумите за бодирање
OPEN_ENDED_ANSWERS = [
    {
        "question": "2. Што му се расипа на Марко еден вторник попладне?",
        "process": "Пронаоѓање на експлицитно наведени информации",
        "points": {
            "1 поен (Точен одговор)": "Механичкиот пес од лимена конзерва / Играчката куче / Му се скрши тркалцето на песот.",
            "0 поени (Неточен одговор)": "Одговор кој не се однесува на текстот (пр. велосипедот, часовникот)."
        }
    },
    {
        "question": "7. Како се промени односот помеѓу Марко и господин Петар?",
        "process": "Интерпретација и интегрирање на идеи",
        "points": {
            "2 поени (Целосно разбирање)": "Ученикот ги опишува двете состојби – и почетната (страв/избегнување) и крајната (пријателство). Пример: На почетокот Марко се плашеше од него, но на крајот се спријателија.",
            "1 поен (Делумно разбирање)": "Ученикот опишува само една состојба. Пример: На крајот тие станаа многу добри пријатели.",
            "0 поени (Нема разбирање)": "Неточен или неповрзан одговор. Пример: Марко му ги кршеше работите."
        }
    },
    {
        "question": "11. Дали насловот „Тајната на стариот часовничар“ е соодветен за овој расказ? Објасни.",
        "process": "Испитување и вреднување на содржината",
        "points": {
            "2 поени (Изграден став со доказ)": "Одговорот содржи став (Да/Не) и валидно објаснување базирано на содржината. Пример: Да, бидејќи тајната не била само занаетот, туку неговото добро срце.",
            "1 поен (Став со површно објаснување)": "Одговара со „Да“ или „Не“, но дава многу штуро објаснување. Пример: Да, затоа што тој е часовничар.",
            "0 поени (Неприфатлив одговор)": "Дава само став („Да“ или „Не“) без никакво објаснување."
        }
    }
]

# ==========================================
# 2. ФУНКЦИИ ЗА ФОРМАТИРАЊЕ
# ==========================================

def setup_document():
    doc = Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12) # Помал фонт за прирачникот за наставници (стандард 12pt)
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    return doc

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Arial'
    run.font.size = Pt(16 if level == 1 else 14)
    run.font.color.rgb = RGBColor(0, 51, 102) # Темно сина боја за прегледност
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT

# ==========================================
# 3. ГЕНЕРИРАЊЕ НА ДОКУМЕНТОТ
# ==========================================

doc = setup_document()

# Наслов
add_heading(doc, "УПАТСТВО ЗА ОЦЕНУВАЊЕ (Клуч за наставникот)", level=1)
doc.add_paragraph(f"Текст: „{NASLOV_NA_TEKST}“\n").alignment = WD_ALIGN_PARAGRAPH.CENTER

# Дел 1: Повеќечлен избор
add_heading(doc, "Прашања со повеќечлен избор (MCQ) - 1 поен за точен одговор", level=2)
for num, answer, process in MCQ_ANSWERS:
    p = doc.add_paragraph()
    p.add_run(f"Прашање {num}: ").bold = True
    p.add_run(f"{answer} ")
    process_run = p.add_run(f"(Процес: {process})")
    process_run.font.italic = True
    process_run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph("\n" + "="*50 + "\n")

# Дел 2: Отворени прашања
add_heading(doc, "Прашања од отворен тип (Конструирани одговори)", level=2)
for item in OPEN_ENDED_ANSWERS:
    # Прашање
    q_para = doc.add_paragraph()
    q_para.add_run(item["question"]).bold = True
    
    # Когнитивен процес
    proc_para = doc.add_paragraph()
    proc_run = proc_para.add_run(f"(Процес: {item['process']})")
    proc_run.font.italic = True
    proc_run.font.color.rgb = RGBColor(100, 100, 100)
    proc_para.paragraph_format.space_after = Pt(2)

    # Критериуми за бодирање
    for point_title, description in item["points"].items():
        point_para = doc.add_paragraph()
        point_para.paragraph_format.left_indent = Inches(0.5)
        point_para.add_run(f"• {point_title}: ").bold = True
        point_para.add_run(description)

    doc.add_paragraph() # Празен ред меѓу прашањата

# Зачувување
filename = f"PIRLS_Kluc_za_Nastavnik_{NASLOV_NA_TEKST.replace(' ', '_')}.docx"
doc.save(filename)
print(f"Документот за наставникот е успешно генериран: {filename}")
