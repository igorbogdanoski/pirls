import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING

NASLOV_NA_TEKST = "Цвеќиња на покривот"

MCQ_ANSWERS = [
    ("1", "Б", "Изведување непосредни заклучоци"),
    ("2", "В", "Фокусирање и пронаоѓање експлицитни информации"),
    ("4", "Б", "Фокусирање и пронаоѓање експлицитни информации"),
    ("6", "В", "Изведување непосредни заклучоци"),
    ("9", "Б", "Фокусирање и пронаоѓање експлицитни информации"),
    ("11", "Б", "Испитување и вреднување на јазични елементи"),
    ("12", "В", "Интерпретација и интегрирање на идеи"),
    ("13", "Б", "Испитување и вреднување на содржината")
]

OPEN_ENDED_ANSWERS = [
    {
        "question": "2. Наведи две работи кои го правеле изгледот на нејзината селска куќа необичен и убав.",
        "process": "Фокусирање и пронаоѓање експлицитни информации",
        "points": {
            "2 бода (Целосно разбирање)": "Наведени се двете карактеристики: 1) покрив покриен со трева (и цвеќиња) и 2) две мали прозорчиња.",
            "1 бод (Делумно разбирање)": "Наведена е само една од овие карактеристики.",
            "0 бодови (Нема разбирање)": "Неточен одговор."
        }
    },
    {
        "question": "3. Зошто Баба Гун не се чувствувала осамена додека живеела во селото?",
        "process": "Изведување непосредни заклучоци",
        "points": {
            "1 бод (Прифатлив одговор)": "Покажува разбирање дека друштво ѝ правеле нејзините животни. Пример: „Затоа што имала многу животни со кои се забавувала.“",
            "0 бодови (Неприфатлив одговор)": "Неточен одговор."
        }
    },
    {
        "question": "5. Наведи две причини зошто станот не ѝ се допаѓал.",
        "process": "Фокусирање и пронаоѓање експлицитни информации",
        "points": {
            "2 бода (Целосно разбирање)": "Наведени се две причини: 1) Ѕидовите биле мазни и бели, 2) Прозорците биле премногу големи.",
            "1 бод (Делумно разбирање)": "Наведена е само една причина.",
            "0 бодови (Нема разбирање)": "Неточен одговор (пр. „беше лош“)."
        }
    },
    {
        "question": "7. Како реагирала Баба Гун на зборовите на раскажувачот?",
        "process": "Интерпретација и интегрирање на идеи",
        "points": {
            "2 бода (Целосно разбирање)": "Таа не сфатила дека е шега, добила идеја и следниот ден отишла со камион да ги земе животните.",
            "1 бод (Делумно разбирање)": "Само пишува дека ги донела животните, без да го опише моментот на инспирација (очите ѝ светнале).",
            "0 бодови (Нема разбирање)": "Неточен одговор."
        }
    },
    {
        "question": "8. Која зграда ја избра и зошто?",
        "process": "Интерпретација и интегрирање идеи / Визуелна евалуација",
        "points": {
            "2 бода (Целосно разбирање)": "Го идентификува точниот избор (Зграда 2 / Зграда Б) И дава доказ од текстот (бидејќи таа насадила трева и ги ставила животните на покривот).",
            "1 бод (Делумно разбирање)": "Го идентификува точниот избор, но објаснувањето е нејасно или недоволно (пр. „Затоа што е убава“).",
            "0 бодови (Нема разбирање)": "Погрешен избор на зграда или нема објаснување."
        }
    },
    {
        "question": "10. Спореди го животот на Баба Гун во селото и во градот. Што е исто?",
        "process": "Интерпретација и интегрирање на идеи",
        "points": {
            "2 бода (Целосно разбирање)": "Поврзува дека во двата случаи таа живее опкружена со своите животни и со природа (трева) и затоа е среќна.",
            "1 бод (Делумно разбирање)": "Поврзува само еден елемент (пр. „ги има истите животни“).",
            "0 бодови (Нема разбирање)": "Неточен одговор."
        }
    },
    {
        "question": "14. Дали мислиш дека насловот „Цвеќиња на покривот“ е соодветен?",
        "process": "Испитување и вреднување на содржината",
        "points": {
            "2 бода (Целосно разбирање)": "Дава став (Да) и го објаснува: цвеќињата ја правеле среќна во селото, а кога ги насадила на покривот во градот, ѝ го вратиле чувството на дом.",
            "1 бод (Делумно разбирање)": "Дава став (Да) со буквално објаснување (пр. „Да, затоа што насади цвеќиња на покривот“).",
            "0 бодови (Незадоволително разбирање)": "Недавање објаснување или нелогичен одговор."
        }
    }
]

def setup_document():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    return doc

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Arial'
    run.font.size = Pt(16 if level == 1 else 14)
    run.font.color.rgb = RGBColor(0, 51, 102)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT

doc = setup_document()

add_heading(doc, "УПАТСТВО ЗА ОЦЕНУВАЊЕ (Клуч за наставникот)", level=1)
doc.add_paragraph(f"Текст: „{NASLOV_NA_TEKST}“\n").alignment = WD_ALIGN_PARAGRAPH.CENTER

add_heading(doc, "Прашања со повеќечлен избор (MCQ) - 1 бод за точен одговор", level=2)
for num, answer, process in MCQ_ANSWERS:
    p = doc.add_paragraph()
    p.add_run(f"Прашање {num}: ").bold = True
    p.add_run(f"{answer} ")
    process_run = p.add_run(f"(Процес: {process})")
    process_run.font.italic = True
    process_run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph("\n" + "="*50 + "\n")

add_heading(doc, "Прашања од отворен тип (Конструирани одговори)", level=2)
for item in OPEN_ENDED_ANSWERS:
    q_para = doc.add_paragraph()
    q_para.add_run(item["question"]).bold = True
    
    proc_para = doc.add_paragraph()
    proc_run = proc_para.add_run(f"(Процес: {item['process']})")
    proc_run.font.italic = True
    proc_run.font.color.rgb = RGBColor(100, 100, 100)
    proc_para.paragraph_format.space_after = Pt(2)

    for point_title, description in item["points"].items():
        point_para = doc.add_paragraph()
        point_para.paragraph_format.left_indent = Inches(0.5)
        point_para.add_run(f"• {point_title}: ").bold = True
        point_para.add_run(description)

    doc.add_paragraph()

filename = "PIRLS_Kluc_za_Nastavnik_Baba_Gun.docx"
doc.save(filename)
print(f"Документот за наставникот е успешно генериран: {filename}")
