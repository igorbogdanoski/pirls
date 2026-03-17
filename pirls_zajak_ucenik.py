from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING

def setup_document():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(14)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    style.paragraph_format.space_after = Pt(12)
    return doc

def add_title(doc, text):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(16)
    run.bold = True

def add_paragraph(doc, text):
    doc.add_paragraph(text).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def add_image_prompt(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(f"[ МЕСТО ЗА ИЛУСТРАЦИЈА: {text} ]")
    run.font.size = Pt(12)
    run.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_question_mcq(doc, number, question, options):
    doc.add_paragraph(f"{number}. {question}").runs[0].bold = True
    for option in options:
        doc.add_paragraph(f"O  {option}").paragraph_format.left_indent = Inches(0.5)

def add_question_open(doc, number, question, points, lines=2):
    p = doc.add_paragraph()
    p.add_run(f"{number}. {question}").bold = True
    p.add_run(f" [✎ {points} поен/и]").italic = True
    for _ in range(lines):
        doc.add_paragraph("_" * 60)

doc = setup_document()

# НАСОКИ
add_title(doc, "НАСОКИ ЗА УЧЕНИКОТ")
add_paragraph(doc, "Во овој тест ќе прочиташ еден расказ и ќе одговориш на прашања за него. Дади сè од себе да ги одговориш сите прашања.")
doc.add_page_break()

# ТЕКСТ
add_title(doc, "ЗАЈАКОТ ГО НАЈАВУВА ЗЕМЈОТРЕСОТ")
add_paragraph(doc, "Си бил некогаш некој зајак кој секогаш бил загрижен. „Оф, леле“ мрморел тој по цел ден. Најмногу бил загрижен за тоа дека може да има земјотрес. Посебно бил загрижен едно утро кога еден огромен плод паднал од дрво - ТРАС - од кој целата земја се стресла. „Земјотрес!” - извикал тој и почнал да бега.")
add_image_prompt(doc, "Зајакот потскокнува во паника додека плодот паѓа на земја.")
add_paragraph(doc, "Десет илјади зајаци почнале да трчаат по него преку полињата и планините. На врвот на планината ги сретнал еден лав. Лавот прашал кој го слушнал земјотресот, па зајакот му објаснил што се случило. Лавот го повикал зајакот да му покаже каде се случила несреќата.")
add_paragraph(doc, "Кога стигнале назад, лавот го зел плодот и повторно го фрлил на земја. ТРАС! Зајакот пак се исплашил, но лавот почнал да се смее. Зајакот сфатил дека тоа не било земјотрес. Лавот љубезно му објаснил дека сите понекогаш се плашиме од работи кои не ги разбираме.")
add_image_prompt(doc, "Лавот му ја објаснува вистината на зајакот на пријателски начин.")
doc.add_page_break()

# ПРАШАЊА
add_title(doc, "ПРАШАЊА ЗА ТЕКСТОТ")
add_question_mcq(doc, "1", "Која била најголемата грижа на зајакот?", ["А) лавот", "Б) ударот на плодот", "В) земјотресот", "Г) паднатиот плод"])
add_question_mcq(doc, "2", "Од што се затресла целата земја на почетокот?", ["А) од земјотрес", "Б) од еден огромен плод", "В) од разбеганите зајаци", "Г) од паѓањето на едно дрво"])
add_question_open(doc, "3", "Најди и препиши два збора што покажуваат дека работите се случувале брзо.", 1, lines=2)
add_question_open(doc, "4", "Каде сакал лавот да го однесе зајакот?", 1, lines=2)
add_question_mcq(doc, "5", "Зошто лавот го испуштил плодот повторно на земја?", ["А) да го натера да бега", "Б) да му помогне", "В) да му покаже што навистина се случило", "Г) да го насмее"])
add_question_mcq(doc, "6", "Како се чувствувал зајакот откако сфатил дека е тоа само плод?", ["А) налутено", "Б) разочарано", "В) глупаво", "Г) загрижено"])
add_question_open(doc, "7", "Напиши два начини на кои лавот го утешил зајакот на крајот.", 2, lines=2)
add_question_open(doc, "8", "Мислиш ли дека на лавот му се допаѓал зајакот? Што го покажува тоа?", 2, lines=3)
add_question_open(doc, "9", "Како се менуваат чувствата на зајакот низ приказната? Објасни.", 3, lines=3)
add_question_open(doc, "10", "Опиши во што се разликуваат лавот и зајакот според нивните постапки.", 3, lines=3)
add_question_mcq(doc, "11", "Која е главната порака во приказната?", ["А) Бегај од неволја", "Б) Провери ги фактите пред паника", "В) Не му верувај на лав", "Г) Зајаците се брзи"])
add_question_mcq(doc, "12", "Што видел првиот зајак кога погледнал назад од планината?", ["А) Земјотрес", "Б) Лав", "В) Стадо зајаци", "Г) Плод"])
add_question_open(doc, "13", "Зошто зајакот се согласил да го води лавот иако се плашел?", 1, lines=2)
add_question_open(doc, "14", "Што значи споредбата дека зајаците тропале „како гром“?", 2, lines=2)
add_question_open(doc, "15", "Дали лавот постапил паметно што го вратил зајакот назад? Зошто?", 2, lines=3)

add_paragraph(doc, "\n🛑 СТОП. КРАЈ НА ОВОЈ ДЕЛ.")
doc.save("PIRLS_Test_Zajak_Ucenik.docx")
