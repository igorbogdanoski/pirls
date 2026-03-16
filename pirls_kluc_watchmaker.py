from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING

NASLOV_NA_TEKST = "Тајната на стариот часовничар"

MCQ_ANSWERS = [
    ("1", "Б", "Пронаоѓање експлицитни информации"),
    ("3", "В", "Пронаоѓање експлицитни информации"),
    ("5", "В", "Интерпретација и интегрирање идеи"),
    ("7", "Б", "Испитување и вреднување на јазични елементи"),
    ("9", "А", "Интерпретација и интегрирање идеи"),
    ("11", "Б", "Испитување и вреднување на содржината"),
    ("13", "Б", "Интерпретација и интегрирање идеи")
]

OPEN_ENDED_ANSWERS = [
    {"question": "2. Што му се расипа на Марко?", "points": {"1 бод": "Мал механички пес од лимена конзерва."}},
    {"question": "4. Зошто Марко брзо побегна?", "points": {"1 бод": "Се плашел од господин Петар кој бил намуртен и ретко зборувал."}},
    {"question": "6. Опиши ја работилницата според звуците.", "points": {"1 бод": "Тивко чукање (тик-так, чук-чук)."}},
    {"question": "8. Промена на односот.", "points": {"2 бода": "На почетокот се избегнувале и плашеле, на крајот станале пријатели и соработници."}},
    {"question": "10. Зошто господин Петар бил намуртен?", "points": {"2 бода": "Бил осамен човек чиј свет бил „скршен“."}},
    {"question": "12. Две работи што ги научил Марко.", "points": {"1 бод": "Користење пинцети, како работат запчаниците, вредноста на трпението."}},
    {"question": "14. Дали насловот е соодветен?", "points": {"2 бода": "Да, бидејќи ја открива тајната на душата на часовничарот и неговата добрина."}},
    {"question": "15. Доказ од последниот пасус.", "points": {"1 бод": "„Постојано се слушаше детска смеа“."}}
]

def setup_document():
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    return doc

doc = setup_document()
p = doc.add_paragraph("УПАТСТВО ЗА ОЦЕНУВАЊЕ")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].bold = True

for num, ans, proc in MCQ_ANSWERS:
    doc.add_paragraph(f"Прашање {num}: {ans} (Процес: {proc})")

for item in OPEN_ENDED_ANSWERS:
    doc.add_paragraph(f"\n{item['question']}").runs[0].bold = True
    for p_val, desc in item['points'].items():
        doc.add_paragraph(f"• {p_val}: {desc}")

doc.save("PIRLS_Kluc_za_Nastavnik_Watchmaker.docx")
print("Документот за наставникот е успешно генериран!")
