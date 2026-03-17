from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING

def setup_document():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(14)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    style.paragraph_format.space_after = Pt(12)
    return doc

def add_title(doc, text, level=1):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(16 if level == 1 else 14)
    run.bold = True

def add_paragraph(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def add_image_prompt(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(f"[ 🎨 NanoBanana Prompt: {text} ]")
    run.font.size = Pt(11)
    run.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_question_mcq(doc, number, question, options):
    p = doc.add_paragraph()
    p.add_run(f"{number}. {question}").bold = True
    for option in options:
        opt = doc.add_paragraph(f"O  {option}")
        opt.paragraph_format.left_indent = Inches(0.5)

def add_question_open(doc, number, question, points, lines=2):
    p = doc.add_paragraph()
    p.add_run(f"{number}. {question}").bold = True
    p.add_run(f" [✎ {points} поен/и]").italic = True
    for _ in range(lines):
        doc.add_paragraph("_" * 60)

doc = setup_document()

# НАСОКИ
add_title(doc, "ДЕЛ 1: ТЕСТ КНИШКА ЗА УЧЕНИКОТ (ПРОШИРЕНА)")
add_title(doc, "НАСОКИ ЗА УЧЕНИКОТ", level=2)
add_paragraph(doc, "Во овој тест ќе прочиташ една приказна и ќе одговориш на неколку прашања за неа. Дади сè од себе да ги одговориш сите прашања.")
doc.add_page_break()

# ТЕКСТ
add_title(doc, "ЛЕТАЈ, ОРЛЕ, ЛЕТАЈ")
doc.add_paragraph("(Африканска приказна прераскажана од Кристофер Грегоровски)").alignment = WD_ALIGN_PARAGRAPH.CENTER

add_paragraph(doc, "Еден ден, еден фармер излегол да го бара изгубеното теле. Тој отишол во долината и барал покрај реката, меѓу трските, зад карпите и во надојдената вода. Се искачил на падините на високата планина со карпести гребени. Погледнал зад една голема карпа, случајно телето да не се стуткало тука за да се засолни од бурата. И, тогаш запрел. Таму, на работ на карпата, здогледал нешто неочекувано. Мало орле што излегло од јајцето ден-два претходно, а силната бура го однела од гнездото.")
add_paragraph(doc, "Тој посегнал и го зел орлето во двете раце. Решил да го земе дома и да се грижи за него. Речиси пристигнал дома, кога децата со истрчале да го пречекаат. „Телето си се врати само!“ извикале тие.")

add_image_prompt(doc, "A kind farmer holding a tiny, fluffy eaglet in his big hands on a high mountain cliff during a misty morning. In the style of a clean children's picture book, pastel colors, white background bleed --ar 16:9")

add_paragraph(doc, "Фармерот бил многу задоволен. Тој му го покажал орлето на семејството, а потоа внимателно го сместил во кокошарникот меѓу квачките и пилињата. „Орелот е кралот на птиците“, рекол, „но, ние ќе го научиме да биде кокошка“.")
add_paragraph(doc, "И така, орлето живеело со кокошките, и правело сѐ како што правеле тие. Како што растело, почнало да изгледа сосема различно од сите кокошки. Еден ден, еден пријател им дошол во посета. Пријателот ја видел птицата меѓу кокошките. „Еј! Тоа не е кокошка. Тоа е орел!“ Фармерот се насмеал и рекол „Се разбира дека е кокошка. Погледни – оди како кокошка, јаде како кокошка. Размислува како кокошка.“")

add_image_prompt(doc, "A large, majestic eagle with brown feathers sitting awkwardly among a group of small white chickens in a sunny farmyard. In the style of a clean children's picture book, pastel colors, white background bleed --ar 16:9")

add_paragraph(doc, "Пријателот не бил убеден. Ја подигнал птицата над главата и рекол: „Ти не си кокошка, ти си орел. Ти не припаѓаш долу на земјата, туку на небото. Летај, орле, летај!“ Птицата ги раширила своите крилја, ги видела кокошките како јадат и скокнула долу да колве храна со нив. „Ти реков дека е кокошка“, рекол фармерот.")
add_paragraph(doc, "Наредниот ден, уште рано наутро, пријателот повторно дошол. „Дај ми уште една шанса со птицата. Дојди со мене на планината каде што ја најде.“ Двајцата мажи тргнале во темнината. „За да може орелот да го види сонцето како изгрева и да го следи до небото каде припаѓа.“ Конечно, пријателот ја донел птицата на самиот раб на планината.")

add_image_prompt(doc, "Two men standing on a high mountain peak at dawn. One man is holding a large eagle towards a magnificent golden and pink sunrise. In the style of a clean children's picture book, pastel colors, white background bleed --ar 16:9")

add_paragraph(doc, "„Погледни го сонцето, Орле. И кога тоа ќе изгрее, летни со него. Ти припаѓаш на небото, а не на земјата.“ Во тој миг, првите сончеви зраци се појавија и светот беше облеан со светлина. Орелот ја протегна главата, ги рашири крилјата нанадвор. Потоа, чувствувајќи го струењето на воздухот, големиот орел се наведна напред и се подигна нагоре, повисоко и повисоко, и се изгуби од погледот во светлината на сонцето, за никогаш повеќе да не биде виден како живее меѓу кокошки.")

doc.add_page_break()

# ПРАШАЊА
add_title(doc, "ПРАШАЊА ЗА ТЕКСТОТ")

add_question_mcq(doc, "1", "Што тргнал да бара фармерот на почетокот на приказната?", ["А. теле", "Б. сточари", "В. карпести гребени", "Г. орле"])
add_question_open(doc, "2", "Каде точно фармерот го нашол малото орле?", 1, lines=2)
add_question_mcq(doc, "3", "Што во првиот дел од приказната ни покажува дека фармерот бил нежен со птицата?", ["А. Го носел орлето во двете раце.", "Б. Го донел орлето кај своето семејство.", "В. Го вратил орлето во гнездото.", "Г. Го барал орлето покрај реката."])
add_question_open(doc, "4", "Што сакал фармерот да го научи орелот откако го сместил во кокошарникот?", 1, lines=2)
add_question_open(doc, "5", "Во текот на првата посета на пријателот, орлето се однесувало како кокошка. Наведете два примери за тоа.", 2, lines=2)
add_question_mcq(doc, "6", "Кога пријателот на фармерот прво го видел орелот, како се обидел да го натера да лета?", ["А. Го подигнал над главата.", "Б. Го ставил на земјата.", "В. Го фрлил во воздух.", "Г. Го однел во планината."])
add_question_mcq(doc, "7", "Зошто орелот не полетал првиот пат кога пријателот го подигнал во дворот?", ["А. Бидејќи бил премногу тежок.", "Б. Бидејќи ги видел кокошките како јадат и сакал да им се придружи.", "В. Бидејќи се плашел од децата на фармерот.", "Г. Бидејќи му биле повредени крилјата."])
add_question_open(doc, "8", "Објаснете што сакал да каже пријателот со зборовите: „Ти не припаѓаш на земјата, туку на небото“.", 2, lines=2)
add_question_open(doc, "9", "Зошто фармерот „громогласно се насмеал“ при првата посета на неговиот пријател?", 1, lines=2)
add_question_open(doc, "10", "Зошто пријателот го однел орелот на високата планина во зори? Наведете две причини.", 2, lines=2)
add_question_open(doc, "11", "Најдете ги и препишете го зборовите од текстот што опишуваат како изгледале облаците во зората.", 1, lines=2)
add_question_mcq(doc, "12", "Зошто сонцето што изгрева е клучно за орелот на крајот од приказната?", ["А. Го разбудило неговиот природен инстинкт да лета.", "Б. Му ги осветлило планините за да не падне.", "В. Го натерало фармерот да замолчи.", "Г. Му помогнало на пријателот да ја види патеката."])
add_question_open(doc, "13", "Опишете каков бил пријателот на фармерот. Наведете една негова особина и пример од неговите постапки.", 2, lines=3)
add_question_mcq(doc, "14", "Која е главната поука на оваа приказна?", ["А. Секогаш треба да го слушаме фармерот.", "Б. Кокошките и орлите можат да бидат добри пријатели.", "В. Секој треба да ја открие својата вистинска природа и да си ги оствари можностите.", "Г. Планините се опасни места за луѓето."])
add_question_open(doc, "15", "Дали мислите дека насловот „Летај, Орле, летај“ е соодветен? Објаснете го вашиот одговор користејќи го она што го прочитавте.", 2, lines=3)

add_paragraph(doc, "\n🛑 СТОП. КРАЈ НА ТЕСТОТ.")
doc.save("PIRLS_Test_Ucenik_Eagle_15Q.docx")
print("Документот за ученикот (15 прашања) е генериран.")
