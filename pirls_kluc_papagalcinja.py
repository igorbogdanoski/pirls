from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_teacher_guide_papagalcinja():
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)

    title = doc.add_paragraph("ДЕЛ 2: КЛУЧ ЗА НАСТАВНИКОТ (УПАТСТВО ЗА ОЦЕНУВАЊЕ) - 15 Прашања")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True

    doc.add_paragraph("Текст: Ноќите на морските папагалчиња\n")

    mcq_answers = [
        ("1", "Г"), ("2", "В"), ("3", "В"), ("4", "Б"), ("5", "Б"), ("6", "А"), ("9", "Г"), ("11", "Г")
    ]

    doc.add_paragraph("Прашања со повеќечлен избор (1 поен за точен одговор):").runs[0].bold = True
    for num, ans in mcq_answers:
        doc.add_paragraph(f"Прашање {num}: {ans}")

    doc.add_paragraph("\nПрашања од отворен тип:").runs[0].bold = True
    
    open_guide = [
        ("7", "Ламба", "1 бод: Му помага на детето да види во мракот каде се скриле птиците."),
        ("8", "Кутии", "1 бод: Служат како безбедно место за транспорт на птиците."),
        ("10", "Ден/Ноќ", "2 бода: Дење птиците го гледаат морето и не се збунети од светлата на селото."),
        ("12", "Чувства", "2 бода: Гордост/Среќа (поради спасувањето) и Тага/Носталгија (поради разделбата)."),
        ("13", "Личен став", "1 бод: Да/Не со објаснување поврзано со грижата за птиците."),
        ("14", "Кловнови", "1 бод: Поради шарените клунови или несмасното одење/летање."),
        ("15", "Натпревар", "2 бода: Децата го сфаќаат спасувањето како забавна и возбудлива мисија.")
    ]

    for num, name, criteria in open_guide:
        p = doc.add_paragraph()
        p.add_run(f"Прашање {num} ({name}): ").bold = True
        doc.add_paragraph(criteria)

    doc.save("PIRLS_Kluc_za_Nastavnik_Papagalcinja.docx")

create_teacher_guide_papagalcinja()
