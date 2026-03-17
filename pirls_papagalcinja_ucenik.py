from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING

def setup_document():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(14)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    style.paragraph_format.space_after = Pt(12)
    return doc

def add_title(doc, text, level=1):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(16 if level == 1 else 14)
    run.bold = True

def add_paragraph(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def add_image_prompt(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(f"[ 🎨 NanoBanana Prompt: {text} ]")
    run.font.size = Pt(11)
    run.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_question_mcq(doc, number, question, options):
    p = doc.add_paragraph()
    p.add_run(f"{number}. {question}").bold = True
    for option in options:
        opt = doc.add_paragraph(f"O  {option}")
        opt.paragraph_format.left_indent = Inches(0.5)

def add_question_open(doc, number, question, lines=2):
    p = doc.add_paragraph()
    p.add_run(f"{number}. {question}").bold = True
    for _ in range(lines):
        doc.add_paragraph("_" * 60)

doc = setup_document()

# НАСОКИ
add_title(doc, "ДЕЛ 1: ТЕСТ КНИШКА ЗА УЧЕНИКОТ")
add_title(doc, "НАСОКИ ЗА УЧЕНИКОТ", level=2)
add_paragraph(doc, "Во овој тест ќе прочиташ еден информативен текст и ќе одговориш на неколку прашања за она што си го прочитал/а. Дади сè од себе да ги одговориш сите прашања.")
doc.add_page_break()

# ТЕКСТ
add_title(doc, "НОЌИТЕ НА МОРСКИТЕ ПАПАГАЛЧИЊА")
doc.add_paragraph("од Брус МекМилан").alignment = WD_ALIGN_PARAGRAPH.CENTER

add_paragraph(doc, "Хана живее на островот Хејмаеј. Секој ден пребарува по небото. Додека гледа од високо, од еден стрмен гребен надвиснат над морето, го здогледува својот прв морски папагал за таа сезона. Си шепнува за себе - „Лунди“, што на исландски значи морски папагал.")
add_paragraph(doc, "Набргу небото е ишарано со нив како со точки - морски папагали, насекаде има морски папагали. Од својот зимски престој на море тие се враќаат на островот на Хана и на блиските ненаселени острови за да снесат јајца и да ги израснат своите папагалчиња. Овие „кловнови на морето“ секоја година се враќаат во истите гнезда. Тоа е единствен пат кога доаѓаат на брегот.")

add_image_prompt(doc, "Children's book illustration, watercolor style. A creative 10-year-old girl standing on a high cliff in Iceland, looking at thousands of puffins flying in the sky over the blue ocean. --ar 16:9")

add_paragraph(doc, "Секоја година, црно-бели птици со портокалови клунови го посетуваат исландскиот остров Хејмаеј. Овие птици ги викаат морски папагали. Познати се и како „кловнови на морето“ поради нивните светли клунови и невешти движења. Морските папагали се чудни летачи при полетување и слетување бидејќи имаат крупни тела и куси крилја.")
add_paragraph(doc, "Хана и нејзините другарчиња се искачуваат на стрмните гребени за да ги набљудуваат птиците. Гледаат како паровите птици се потчукнуваат една со друга со клуновите. Секој пар што го набљудуваат наскоро ќе се грижи за јајце сместено длабоко во стрмниот гребен. Кога од птичјите јајца ќе се изведат малите пиленца, родителите им носат риба во гнездото за да ги нахранат. Секое пиленце ќе израсне во младо морско папагалче.")

add_image_prompt(doc, "Close up illustration, watercolor style. A pair of puffins touching their colorful orange beaks together on a grassy cliff. Soft evening light. --ar 3:2")

add_paragraph(doc, "Цело лето возрасните морски папагали ловат риби и се грижат за своите пиленца. До август, цвеќињата ќе ги покријат гнездата. Кога цвеќињата целосно ќе расцветаат, Хана знае дека завршило чекањето на ноќите на морските папагалчиња. Скриените пиленца веќе израснале во млади морски папагалчиња. Сега е време Хана и нејзините другарчиња да ги извадат своите картонски кутии и батериски ламби за ноќите на морските папагалчиња.")
add_paragraph(doc, "Во темнината на ноќта, морските папагалчиња ги напуштаат своите гнезда за да го направат својот прв лет. Најголем дел од птиците успешно слетуваат во морето под нив. Но некои од нив се збунуваат од селските светла – веројатно мислат дека тоа се зраците на месечината што се одбиваат од водата. Стотици млади морски папагалчиња секоја ноќ треснуваат на земја во селото. Неспособни да полетаат од рамното земјиште, тие трчаат наоколу и се обидуваат да се скријат.")

add_image_prompt(doc, "Night scene, watercolor style. A young girl with a flashlight and a cardboard box rescuing a small puffin on a village street at night. --ar 16:9")

add_paragraph(doc, "Хана и нејзините пријатели ќе ја поминат секоја ноќ барајќи ги беспомошните морски папагалчиња кои не успеале да стигнат до водата. Децата мора први да ги најдат пред селските мачки или автомобилите. Веќе две недели сите деца од Хејмаеј спијат до доцна наутро за да можат да останат надвор до доцна во ноќта. Тие спасуваат илјадници млади папагалчиња.")
add_paragraph(doc, "Време е морските папагалчиња да бидат пуштени на слобода. Хана прва ослободува едно од нив. Го држи високо за да може папагалчето да се навикне да мавта со крилјата. Потоа, замавнува со него нагоре во воздухот и го фрла кон водата покрај стрмниот гребен. Морското папагалче мавта со крилјата и успешно слетува, плеснувајќи се во водата. Хана се збогува со нив до следната пролет. „До видување, до видување!“. ")

add_image_prompt(doc, "Beautiful morning scene, watercolor. A girl on a beach throwing a puffin into the air towards the sea waves. The sun is rising. --ar 16:9")

doc.add_page_break()

# ПРАШАЊА
add_title(doc, "ПРАШАЊА ЗА ТЕКСТОТ")

questions = [
    (1, "Зошто морските папагали се невешти при полетувањето и слетувањето?", ["Живеат во ледена земја.", "Скоро никогаш не излегуваат на брегот.", "Времето го поминуваат на високи стрмни гребени.", "Имаат крупни тела и куси крилја."]),
    (2, "Каде ја минуваат зимата морските папагали?", ["во стрмните гребени", "на плажата", "на море", "на мразот"]),
    (3, "Зошто морските папагали доаѓаат на островот?", ["да бидат спасени", "да бараат храна", "да снесат јајца", "да научат да летаат"]),
    (4, "Како Хана знае дека е време морските папагалчиња да полетаат?", ["Родителите им носат риба.", "Цвеќињата се целосно расцветани.", "Пиленцата се скриени.", "Летото тукушто почнало."]),
    (5, "Што се случува за време на ноќите на морските папагалчиња?", ["Паровите се потчукнуваат.", "Го прават својот прв лет.", "Се изведуваат од јајцата.", "Доаѓаат од море на брегот."]),
    (6, "Што можеле луѓето од селото да направат за да ги спречат да слетуваат таму?", ["да ги изгаснат светлата", "да ги подготват кутиите", "да ги чуваат мачките дома", "да свртат ламби кон небото"])
]

for num, q, opts in questions:
    add_question_mcq(doc, num, q, opts)

add_question_open(doc, 7, "Објасни како Хана ја користи батериската ламба за да ги спаси морските папагалчиња.", lines=2)
add_question_open(doc, 8, "Објасни како Хана ги користи картонските кутии за да ги спаси морските папагалчиња.", lines=2)
add_question_mcq(doc, 9, "Според текстот, со каква опасност биле соочени морските папагалчиња во селото?", ["да се удават во морето", "да се загубат во гнездата", "да не добиваат риби", "да бидат прегазени од автомобили"])
add_question_open(doc, 10, "Зошто треба да е дење кога децата ги ослободуваат морските папагалчиња? Објасни со информации од текстот.", lines=3)
add_question_mcq(doc, 11, "Што прават морските папагалчиња откако Хана ќе ги ослободи?", ["шетаат по плажата", "полетуваат од гребени", "се кријат во селото", "пливаат во морето"])
add_question_open(doc, 12, "Напиши две чувства кои Хана можеби ги имала по ослободувањето и објасни зошто.", lines=4)
add_question_open(doc, 13, "Дали ти би сакал/а да спасуваш морски папагалчиња со Хана? Објасни зошто.", lines=3)
add_question_open(doc, 14, "Зошто овие птици луѓето често ги нарекуваат „кловнови на морето“?", lines=2)
add_question_open(doc, 15, "Што ни покажува тоа што децата се натпреваруваат во спасувањето?", lines=3)

add_paragraph(doc, "\n🛑 СТОП. КРАЈ НА ТЕСТОТ.")
doc.save("PIRLS_Test_Kniska_Papagalcinja.docx")
print("Документот за ученикот е генериран.")
