from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING

def setup_document():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(14)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    style.paragraph_format.space_after = Pt(12)
    return doc

def add_title(doc, text, level=1):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(16 if level == 1 else 14)
    run.bold = True

def add_paragraph(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def add_image_prompt(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(f"[ 🎨 NanoBanana Prompt: {text} ]")
    run.font.size = Pt(11)
    run.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_question_mcq(doc, number, question, options):
    p = doc.add_paragraph()
    p.add_run(f"{number}. {question}").bold = True
    for option in options:
        opt = doc.add_paragraph(f"O  {option}")
        opt.paragraph_format.left_indent = Inches(0.5)

def add_question_open(doc, number, question, points, lines=2):
    p = doc.add_paragraph()
    p.add_run(f"{number}. {question}").bold = True
    p.add_run(f" [✎ {points} поен/и]").italic = True
    for _ in range(lines):
        doc.add_paragraph("_" * 60)

doc = setup_document()

# НАСОКИ
add_title(doc, "ДЕЛ 1: ТЕСТ КНИШКА ЗА УЧЕНИКОТ")
add_title(doc, "НАСОКИ ЗА УЧЕНИКОТ", level=2)
add_paragraph(doc, "Во овој тест ќе прочиташ една приказна и ќе одговориш на неколку прашања за неа. Дади сè од себе да ги одговориш сите прашања.")
doc.add_page_break()

# ТЕКСТ
add_title(doc, "ПИТА ЗА НЕПРИЈАТЕЛОТ")
doc.add_paragraph("(од Дерек Мунсон)").alignment = WD_ALIGN_PARAGRAPH.CENTER

add_paragraph(doc, "Летото беше совршено сѐ додека Џереми Рос не се всели во куќата веднаш до куќата на мојот најдобар пријател Стенли. Не ми се допаѓаше Џереми. Организираше забава, а мене дури не ме ни покани. Но, мојот најдобар другар Стенли беше поканет. Никогаш сум немал непријател, сѐ додека Џереми не се всели во соседството. Тато му кажа дека кога бил на моја возраст, и тој имал непријатели. Но, знаел начин како да се спаси од нив.")
add_paragraph(doc, "Тато извади еден излитен лист хартија од книгата со рецепти. „Пита за непријателот“, рече задоволно. „Пита за непријателот е најбрзиот познат начин за да се ослободиш од непријателите.“. Тоа ме замисли. Какви грозни работи би ставил во Питата за непријателот? Му однесов глисти и камења на тато, но тој веднаш ми ги врати.")

add_image_prompt(doc, "A creative 10-year-old boy named Tom standing in a kitchen, offering a handful of worms and small stones to his father. The father is smiling and pointing to an old recipe book on the counter. In the style of a clean children's picture book, pastel colors, white background bleed --ar 16:9")

add_paragraph(doc, "Се обидов да замислам колку одвратно мириса Питата за непријателот. Но, мирисаше навистина добро. Влегов внатре да го прашам тато што не е во ред. Питата за непријателот не би смеела да мириса толку добро. Но, тато беше паметен. „Ако мириса лошо, непријателот никогаш не би ја изел“, рече.")
add_paragraph(doc, "Додека питата се ладеше, тато ми кажа што е мојата задача. Тој шепотеше. „За да делува, треба да поминеш еден ден со непријателот. Дури и полошо, мора да бидеш љубезен со него. Не е лесно. Но, тоа е единствениот начин Питата за непријателот да делува.“")

add_image_prompt(doc, "Tom riding his bicycle towards a neighbor's house, looking a bit nervous but determined. The sun is shining brightly. In the background, a cozy suburban street. In the style of a clean children's picture book, pastel colors, white background bleed --ar 16:9")

add_paragraph(doc, "Сѐ што требаше да направам беше да поминам еден ден со Џереми. Малку возевме велосипеди, а потоа ручавме. По ручекот, отиводме кај мене дома. Беше чудно, но се забавував со мојот непријател. Игравме игри додека тато не нѐ повика за вечера. Почнав да си мислам дека треба да ја заборавиме Питата за непријателот. „Тато“, реков, „Навистина е убаво да имаш нов пријател.“")
add_paragraph(doc, "Но по вечерата, тато ја донесе питата. „Вау!“ рече Џереми, гледајќи ја питата. Ме фати паника. Не сакав Џереми да ја изеде Питата за непријателот! Тој ми беше пријател! „Не јади ја!“ Извикав. „Лоша е!“ Виљушката на Џереми запре. Ме погледна чудно. „Ако е толку лоша“, праша Џереми, „тогаш зошто татко ти веќе ја изел половината?“")

add_image_prompt(doc, "Three people sitting at a dinner table: Tom, Jeremy, and the Father. The Father is happily eating a large slice of pie. Jeremy is holding a fork, looking surprised, while Tom looks panicked, trying to stop him. In the style of a clean children's picture book, pastel colors, white background bleed --ar 16:9")

add_paragraph(doc, "Се чинеше безбедно, па и јас каснав мал залак. Беше превкусно! По десертот, Џереми ме покани да одам кај него дома наредниот ден. А што се однесува до Питата за непријателот, сѐ уште не знам како се прави. Но, јас само што го изгубив најдобриот непријател.")

add_image_prompt(doc, "Two young boys, Tom and Jeremy, walking together outside, laughing and talking like best friends. A warm sunset in the background. In the style of a clean children's picture book, pastel colors, white background bleed --ar 16:9")

doc.add_page_break()

# ПРАШАЊА
add_title(doc, "ПРАШАЊА ЗА ТЕКСТОТ")

add_question_mcq(doc, "1", "Кој ја раскажува приказната?", ["А. Џереми", "Б. Тато", "В. Стенли", "Г. Том"])
add_question_open(doc, "2", "На почетокот на приказната, зошто Том мислел дека Џереми му е непријател?", 1, lines=2)
add_question_open(doc, "3", "Напишете една состојка што Том мислел дека треба да биде во Питата за непријателот.", 1, lines=2)
add_question_mcq(doc, "4", "Зошто Том мислел дека летото сепак ќе биде добро додека татко му готвел?", ["А. Сакал да си игра надвор.", "Б. Бил возбуден за планот на тато.", "В. Добил нов пријател."])
add_question_open(doc, "5", "Како се чувствувал Том кога првпат ја мириснал питата? Објаснете зошто.", 2, lines=2)
add_question_open(doc, "6", "Што мислел Том дека ќе му се случи на Џереми откако ќе ја изеде питата? Напишете едно нешто.", 1, lines=2)
add_question_open(doc, "7", "Кои две работи му ги кажал тато на Том да ги направи за питата да делува?", 2, lines=2)
add_question_mcq(doc, "8", "Зошто Том отишол во куќата на Џереми?", ["А. Да го покани на вечера.", "Б. Да му се развика.", "В. Да го викне да си играат."])
add_question_open(doc, "9", "Што го изненадило Том во врска со денот поминат со Џереми?", 1, lines=2)
add_question_mcq(doc, "10", "Зошто на вечерата Том сакал да заборават на Питата за непријателот?", ["А. Бил скржав.", "Б. Не верувал во тато.", "В. Почнал да му се допаѓа Џереми."])
add_question_mcq(doc, "11", "Како се чувствувал Том кога тато му подал парче пита на Џереми?", ["А. Вознемирено", "Б. Задоволно", "В. Изненадено", "Г. Збунето"])
add_question_mcq(doc, "12", "Која била вистинската тајна на тато за Питата за непријателот?", ["А. Тоа била обична пита.", "Б. Имала одвратен вкус.", "В. Била отровна."])
add_question_mcq(doc, "13", "Што ни кажува поканата на Џереми за наредниот ден?", ["А. Сѐ уште се непријатели.", "Б. Сакаат уште пита.", "В. Можат да станат пријатели."])
add_question_open(doc, "14", "Објаснете зошто таткото на Том навистина ја направил Питата за непријателот.", 1, lines=2)
add_question_open(doc, "15", "Каква личност е таткото? Наведете особина и пример од приказната.", 2, lines=3)
add_question_open(doc, "16", "Каква поука може да извлечете од оваа приказна?", 1, lines=2)

add_paragraph(doc, "\n🛑 СТОП. КРАЈ НА ТЕСТОТ.")
doc.save("PIRLS_Test_Pita_Ucenik.docx")
print("Документот за ученикот (Пита) е генериран.")
