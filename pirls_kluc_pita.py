from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_pirls_key():
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)

    title = doc.add_paragraph("УПАТСТВО ЗА ОЦЕНУВАЊЕ (КЛУЧ): Пита за непријателот (16 Прашања)")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True

    doc.add_paragraph("Текст: Детска приказна\n")

    mcq_answers = [
        ("1", "Г"), ("4", "Б"), ("8", "В"), ("10", "В"), ("11", "А"), ("12", "А"), ("13", "В")
    ]

    doc.add_paragraph("Прашања со повеќечлен избор (1 поен за точен одговор):").runs[0].bold = True
    for num, ans in mcq_answers:
        doc.add_paragraph(f"Прашање {num}: {ans}")

    doc.add_paragraph("\nПрашања од отворен тип:").runs[0].bold = True
    
    open_guide = [
        ("2", "Причина", "1 бод: Бидејќи се вселил веднаш до Стенли и не го поканил на забавата."),
        ("3", "Состојка", "1 бод: Глисти / камења."),
        ("5", "Чувства", "2 бода: Изненаден / збунет (1 бод) бидејќи питата мирисала убаво (1 бод)."),
        ("6", "Ефект", "1 бод: Мислел дека ќе му паднат забите / ќе му опадне косата / ќе му се случи нешто лошо."),
        ("7", "Задачи", "2 бода: Да помине еден ден со него и да биде љубезен."),
        ("9", "Изненадување", "1 бод: Дека всушност се забавувал со својот непријател."),
        ("14", "Вистинска намена", "1 бод: За да го натера синот да помине време со Џереми и да се спријателат."),
        ("15", "Особина", "2 бода: Мудар/Паметен (1 бод) - пример со измислената пита за да ги зближи децата (1 бод)."),
        ("16", "Поука", "1 бод: Најдобар начин да се победи непријателот е да се претвори во пријател.")
    ]

    for num, name, criteria in open_guide:
        p = doc.add_paragraph()
        p.add_run(f"Прашање {num} ({name}): ").bold = True
        doc.add_paragraph(criteria)

    doc.save("PIRLS_Key_Pita_Ucenik.docx")
    # print is removed to avoid encoding issues in CMD

create_pirls_key()
