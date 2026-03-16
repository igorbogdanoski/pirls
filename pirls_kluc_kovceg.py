from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING

NASLOV_NA_TEKST = "Тајната на стариот ковчег"

MCQ_ANSWERS = [
    ("1", "Б", "Пронаоѓање експлицитни информации"),
    ("3", "В", "Пронаоѓање експлицитни информации"),
    ("6", "Б", "Донесување директни заклучоци"),
    ("7", "Г", "Пронаоѓање експлицитни информации"),
    ("11", "Б", "Испитување и евалуација на содржина")
]

OPEN_ENDED_ANSWERS = [
    {"question": "2. Две работи кои му недостасуваа.", "points": {"1 поен": "Брз интернет, видеоигри или врсници."}},
    {"question": "4. Како се чувствуваше и зошто.", "points": {"2 поени": "Разочарано, бидејќи очекувал злато, а нашол стари предмети."}},
    {"question": "5. Трите предмети.", "points": {"1 поен": "Фотографија, камче и свирче."}},
    {"question": "8. Споредба на мислењето.", "points": {"2 поени": "Прво мислел дека се ѓубре, потоа сфатил дека се вредни спомени."}},
    {"question": "9. Значење на цитатот.", "points": {"2 поени": "Материјалните работи се минливи, спомените се вечни."}},
    {"question": "10. Каква личност е дедото.", "points": {"2 поени": "Мудар/Грижлив. Пример: Го научи Марко за вистинските вредности."}},
    {"question": "12. Дали насловот е соодветен.", "points": {"2 поени": "Да, бидејќи ковчегот ја криеше тајната на минатото на дедото."}}
]

def setup_document():
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    return doc

doc = setup_document()
p = doc.add_paragraph("УПАТСТВО ЗА ОЦЕНУВАЊЕ")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].bold = True

for num, ans, proc in MCQ_ANSWERS:
    doc.add_paragraph(f"Прашање {num}: {ans} (Процес: {proc})")

for item in OPEN_ENDED_ANSWERS:
    doc.add_paragraph(f"\n{item['question']}").runs[0].bold = True
    for p_val, desc in item['points'].items():
        doc.add_paragraph(f"• {p_val}: {desc}")

doc.save("PIRLS_Kluc_za_Nastavnik_Kovceg.docx")
print("Документот за наставникот е успешно генериран!")
